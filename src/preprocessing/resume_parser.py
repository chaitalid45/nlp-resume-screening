"""
Converts PDF and DOCX resume files into a structured ParsedResume dataclass.
Supports section detection, contact extraction, and multi-format parsing.
"""

from __future__ import annotations

import re
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional



# Data model

@dataclass
class ParsedResume:
    """Structured representation of a parsed resume."""
    raw_text: str
    file_path: str = ""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    skills: list[str] = field(default_factory=list)
    sections: dict[str, str] = field(default_factory=dict)
    page_count: int = 0
    word_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.raw_text.split())


# Parser

class ResumeParser:
    """
    Parses PDF and DOCX resume files into ParsedResume objects.

    Usage:
        parser = ResumeParser()
        result = parser.parse("resume.pdf")
        print(result.email, result.sections.keys())
    """

    # Regex patterns
    EMAIL_RE = re.compile(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"
    )
    PHONE_RE = re.compile(
        r"(?:\+?\d{1,3}[\s\-.]?)?"           # optional country code
        r"(?:\(?\d{2,4}\)?[\s\-.]?)"          # area code
        r"\d{3,4}[\s\-.]?\d{3,4}"            # local number
    )
    LINKEDIN_RE = re.compile(
        r"linkedin\.com/in/[\w\-]+"
    )
    GITHUB_RE = re.compile(
        r"github\.com/[\w\-]+"
    )

    # Section header keywords (order matters — checked top-to-bottom)
    SECTION_KEYWORDS = [
        "summary", "objective", "profile",
        "experience", "work experience", "employment",
        "education", "academic",
        "skills", "technical skills", "core competencies",
        "projects", "personal projects",
        "certifications", "certificates", "licenses",
        "publications", "research",
        "awards", "achievements", "honors",
        "languages", "volunteering", "interests",
    ]

    def parse(self, path: str | Path) -> ParsedResume:
        """Parse a PDF or DOCX file and return a ParsedResume."""
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"Resume file not found: {path}")

        suffix = p.suffix.lower()

        if suffix == ".pdf":
            text, pages = self._read_pdf(p)
        elif suffix in (".docx", ".doc"):
            text, pages = self._read_docx(p), 0
        elif suffix == ".txt":
            text, pages = p.read_text(encoding="utf-8", errors="ignore"), 0
        else:
            raise ValueError(f"Unsupported file type: {suffix}. Use PDF, DOCX, or TXT.")

        resume = ParsedResume(
            raw_text=text,
            file_path=str(p),
            page_count=pages,
        )
        resume.email    = self._extract_email(text)
        resume.phone    = self._extract_phone(text)
        resume.linkedin = self._extract_linkedin(text)
        resume.github   = self._extract_github(text)
        resume.sections = self._split_sections(text)
        resume.name     = self._extract_name(text, p.stem)

        return resume

    # File readers

    def _read_pdf(self, path: Path) -> tuple[str, int]:
        """Extract text from a PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF is required: pip install PyMuPDF")

        doc = fitz.open(str(path))
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages), len(pages)

    def _read_docx(self, path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required: pip install python-docx")

        doc = docx.Document(str(path))
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    # Contact extraction

    def _extract_email(self, text: str) -> Optional[str]:
        m = self.EMAIL_RE.search(text)
        return m.group().lower() if m else None

    def _extract_phone(self, text: str) -> Optional[str]:
        m = self.PHONE_RE.search(text)
        return m.group().strip() if m else None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        m = self.LINKEDIN_RE.search(text)
        return m.group() if m else None

    def _extract_github(self, text: str) -> Optional[str]:
        m = self.GITHUB_RE.search(text)
        return m.group() if m else None

    def _extract_name(self, text: str, fallback: str) -> str:
        """
        Heuristic: the name is usually the first non-empty line
        that looks like a proper name (2-4 capitalised words, no digits).
        """
        name_re = re.compile(r"^([A-Z][a-z]+(?:\s[A-Z][a-z]+){1,3})$")
        for line in text.splitlines()[:10]:
            line = line.strip()
            if name_re.match(line):
                return line
        return fallback.replace("_", " ").replace("-", " ").title()

    # Section splitting

    def _split_sections(self, text: str) -> dict[str, str]:
        """
        Split resume text into labelled sections by detecting headers.
        Returns a dict like {"experience": "...", "skills": "...", ...}
        """
        lines = text.splitlines()
        sections: dict[str, list[str]] = {}
        current_section = "header"
        sections[current_section] = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Check if line is a section header
            detected = self._detect_section_header(stripped)
            if detected:
                current_section = detected
                sections.setdefault(current_section, [])
            else:
                sections[current_section].append(stripped)

        # Join lines and remove empty sections
        return {
            key: "\n".join(lines_)
            for key, lines_ in sections.items()
            if lines_
        }

    def _detect_section_header(self, line: str) -> Optional[str]:
        """
        Return the canonical section name if `line` looks like a header,
        else return None.
        """
        # Remove common decorators: ALL CAPS, trailing colons, underscores
        clean = re.sub(r"[_\-—=:]+", " ", line).strip().lower()

        # Skip lines that are too long to be headers
        if len(clean.split()) > 6:
            return None

        for keyword in self.SECTION_KEYWORDS:
            if keyword in clean:
                # Normalise to a clean key
                return keyword.replace(" ", "_")

        return None